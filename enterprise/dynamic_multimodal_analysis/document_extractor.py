#!/usr/bin/env python3
"""
Document Content Extractor
文檔內容提取器 - 支持多種格式的文檔內容提取
"""

import os
import logging
from typing import Optional, Dict, Any
from docx import Document
import PyPDF2
import zipfile
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

def extract_document_content(file_path: str, file_ext: str) -> str:
    """
    提取文檔內容
    
    Args:
        file_path: 文件路徑
        file_ext: 文件擴展名
        
    Returns:
        提取的文本內容
    """
    try:
        if file_ext in ['doc', 'docx']:
            return extract_word_content(file_path)
        elif file_ext == 'pdf':
            return extract_pdf_content(file_path)
        elif file_ext == 'txt':
            return extract_text_content(file_path)
        else:
            logger.warning(f"不支援的文件格式: {file_ext}")
            return f"不支援的文件格式: {file_ext}"
            
    except Exception as e:
        logger.error(f"文檔內容提取失敗: {e}")
        return f"文檔內容提取失敗: {str(e)}"

def extract_word_content(file_path: str) -> str:
    """提取Word文檔內容"""
    try:
        # 嘗試使用python-docx處理.docx文件
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 提取表格內容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return "\n".join(content)
        
        # 處理.doc文件（較舊格式）
        else:
            return extract_doc_content_legacy(file_path)
            
    except Exception as e:
        logger.error(f"Word文檔處理失敗: {e}")
        return f"Word文檔處理失敗，可能是格式不支援: {str(e)}"

def extract_doc_content_legacy(file_path: str) -> str:
    """提取舊版.doc文件內容"""
    try:
        # 嘗試使用antiword（如果可用）
        import subprocess
        result = subprocess.run(['antiword', file_path], 
                              capture_output=True, text=True, timeout=30)
        if result.returncode == 0:
            return result.stdout
        
        # 如果antiword不可用，嘗試其他方法
        logger.warning("antiword不可用，嘗試其他方法")
        return extract_doc_with_zipfile(file_path)
        
    except Exception as e:
        logger.error(f"舊版DOC文件處理失敗: {e}")
        return "舊版DOC文件需要特殊處理工具，建議轉換為DOCX格式"

def extract_doc_with_zipfile(file_path: str) -> str:
    """嘗試使用zipfile方法提取DOC內容"""
    try:
        # 某些.doc文件實際上是zip格式
        with zipfile.ZipFile(file_path, 'r') as zip_file:
            # 查找document.xml或類似文件
            for file_name in zip_file.namelist():
                if 'document' in file_name.lower() and file_name.endswith('.xml'):
                    with zip_file.open(file_name) as xml_file:
                        tree = ET.parse(xml_file)
                        root = tree.getroot()
                        
                        # 提取文本內容
                        text_content = []
                        for elem in root.iter():
                            if elem.text:
                                text_content.append(elem.text.strip())
                        
                        return "\n".join(filter(None, text_content))
        
        return "無法提取DOC文件內容，建議使用DOCX格式"
        
    except Exception as e:
        logger.error(f"ZIP方法提取DOC失敗: {e}")
        return "DOC文件格式複雜，建議轉換為DOCX格式後重新上傳"

def extract_pdf_content(file_path: str) -> str:
    """提取PDF內容"""
    try:
        content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    content.append(text.strip())
        
        return "\n".join(content)
        
    except Exception as e:
        logger.error(f"PDF處理失敗: {e}")
        return f"PDF處理失敗: {str(e)}"

def extract_text_content(file_path: str) -> str:
    """提取純文本內容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # 嘗試其他編碼
        try:
            with open(file_path, 'r', encoding='gbk') as file:
                return file.read()
        except:
            with open(file_path, 'r', encoding='big5') as file:
                return file.read()
    except Exception as e:
        logger.error(f"文本文件處理失敗: {e}")
        return f"文本文件處理失敗: {str(e)}"

def analyze_document_structure(content: str) -> Dict[str, Any]:
    """分析文檔結構"""
    lines = content.split('\n')
    
    structure = {
        "total_lines": len(lines),
        "non_empty_lines": len([line for line in lines if line.strip()]),
        "sections": [],
        "tables_detected": 0,
        "lists_detected": 0
    }
    
    # 檢測章節標題
    for i, line in enumerate(lines):
        line = line.strip()
        if line:
            # 檢測可能的標題（數字開頭、全大寫等）
            if (line.startswith(('第', '一、', '二、', '三、', '1.', '2.', '3.')) or
                line.isupper() and len(line) < 50):
                structure["sections"].append({
                    "line_number": i + 1,
                    "title": line,
                    "level": detect_heading_level(line)
                })
            
            # 檢測表格
            if '|' in line and line.count('|') >= 2:
                structure["tables_detected"] += 1
            
            # 檢測列表
            if line.startswith(('•', '-', '*', '○')):
                structure["lists_detected"] += 1
    
    return structure

def detect_heading_level(text: str) -> int:
    """檢測標題級別"""
    if text.startswith('第') and '章' in text:
        return 1
    elif text.startswith(('一、', '二、', '三、', '四、', '五、')):
        return 2
    elif text.startswith(('1.', '2.', '3.', '4.', '5.')):
        return 3
    elif text.isupper():
        return 2
    else:
        return 4

def extract_key_information(content: str) -> Dict[str, Any]:
    """提取關鍵信息並提供清楚的說明"""
    import re
    
    key_info = {
        "dates": [],
        "numbers": [],
        "processes": [],
        "departments": [],
        "systems": []
    }
    
    # 提取日期
    date_patterns = [
        r'\d{4}年\d{1,2}月',
        r'\d{1,2}/\d{1,2}/\d{4}',
        r'\d{4}-\d{2}-\d{2}'
    ]
    
    for pattern in date_patterns:
        matches = re.findall(pattern, content)
        key_info["dates"].extend(matches)
    
    # 提取數字和比例，並提供清楚的說明
    number_with_context = []
    
    # 提取百分比及其上下文，提供清楚說明
    percent_matches = re.finditer(r'([^。！？\n]*?)(\d+(?:\.\d+)?%)([^。！？\n]*)', content)
    for match in percent_matches:
        before, percent, after = match.groups()
        # 清理上下文，提取關鍵信息
        context_before = before.strip()[-30:] if before.strip() else ""
        context_after = after.strip()[:30] if after.strip() else ""
        
        # 構建清楚的說明
        if "自動化" in context_before or "自動化" in context_after:
            description = f"自動化比率: {percent}"
        elif "準確率" in context_before or "準確率" in context_after:
            description = f"準確率: {percent}"
        elif "審核" in context_before or "審核" in context_after:
            description = f"審核相關: {percent}"
        elif "處理" in context_before or "處理" in context_after:
            description = f"處理效率: {percent}"
        else:
            # 提取最相關的詞彙
            key_words = re.findall(r'[核保|審核|處理|作業|效率|比率|準確|自動]', context_before + context_after)
            if key_words:
                description = f"{key_words[0]}相關: {percent}"
            else:
                description = f"重要比率: {percent}"
        
        number_with_context.append(description)
    
    # 提取金額及其說明
    money_matches = re.finditer(r'([^。！？\n]*?)(\d+(?:,\d{3})*(?:\.\d+)?(?:元|萬|億))([^。！？\n]*)', content)
    for match in money_matches:
        before, money, after = match.groups()
        context = (before.strip()[-20:] + " " + after.strip()[:20]).strip()
        if "投資" in context or "成本" in context:
            description = f"投資成本: {money}"
        elif "保費" in context or "費用" in context:
            description = f"費用相關: {money}"
        else:
            description = f"金額: {money}"
        number_with_context.append(description)
    
    # 提取人數、件數等及其說明
    count_matches = re.finditer(r'([^。！？\n]*?)(\d+(?:人|件|天|小時|分鐘))([^。！？\n]*)', content)
    for match in count_matches:
        before, count, after = match.groups()
        context = (before.strip()[-20:] + " " + after.strip()[:20]).strip()
        if "人員" in context or "配置" in context:
            description = f"人員配置: {count}"
        elif "處理" in context or "申請" in context:
            description = f"處理量: {count}"
        elif "時間" in context:
            description = f"處理時間: {count}"
        else:
            description = f"數量統計: {count}"
        number_with_context.append(description)
    
    key_info["numbers"] = number_with_context[:10]  # 限制數量
    
    # 提取核心流程並提供清楚說明
    process_descriptions = []
    
    # 查找具體的SOP流程步驟
    sop_step_patterns = [
        r'([^。！？\n]*(?:第[一二三四五六七八九十]+步|步驟\d+|流程\d+)[^。！？\n]*)',
        r'([^。！？\n]*(?:收文|審核|核保|出單|建檔|轉帳)[^。！？\n]*)',
        r'([^。！？\n]*(?:標準作業流程|SOP)[^。！？\n]*)'
    ]
    
    for pattern in sop_step_patterns:
        matches = re.findall(pattern, content)
        for match in matches:
            clean_match = match.strip()
            if len(clean_match) > 10 and len(clean_match) < 100:  # 過濾長度
                # 清理和格式化流程描述
                if "收文" in clean_match:
                    process_descriptions.append(f"文件收文流程: {clean_match}")
                elif "審核" in clean_match or "核保" in clean_match:
                    process_descriptions.append(f"審核核保流程: {clean_match}")
                elif "出單" in clean_match or "建檔" in clean_match:
                    process_descriptions.append(f"出單建檔流程: {clean_match}")
                elif "SOP" in clean_match or "標準作業" in clean_match:
                    process_descriptions.append(f"標準作業程序: {clean_match}")
                else:
                    process_descriptions.append(f"業務流程: {clean_match}")
    
    # 如果沒有找到具體流程，提取重要的業務描述
    if len(process_descriptions) < 3:
        business_patterns = [
            r'([^。！？\n]*(?:核保|保險|理賠|申請)(?:作業|業務|處理)[^。！？\n]*)',
            r'([^。！？\n]*(?:要保|被保險人|受益人)[^。！？\n]*)',
            r'([^。！？\n]*(?:保單|契約|保費)[^。！？\n]*)'
        ]
        
        for pattern in business_patterns:
            matches = re.findall(pattern, content)
            for match in matches[:2]:
                clean_match = match.strip()
                if len(clean_match) > 15 and len(clean_match) < 80:
                    process_descriptions.append(f"業務說明: {clean_match}")
    
    key_info["processes"] = process_descriptions[:5]  # 限制數量
    
    return key_info

