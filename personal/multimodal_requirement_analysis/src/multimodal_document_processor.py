#!/usr/bin/env python3
"""
多模態文檔處理增強模組
為需求分析服務提供文檔OCR、圖像理解和內容提取功能
"""

import os
import sys
import json
import asyncio
import logging
import tempfile
import base64
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from PIL import Image
import fitz  # PyMuPDF
import docx
from docx import Document

# 添加路徑
sys.path.append('/home/ubuntu/enterprise_deployment')

class MultimodalDocumentProcessor:
    """多模態文檔處理器"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.logger = logging.getLogger(self.__class__.__name__)
        
        # 支持的文件類型
        self.supported_formats = {
            'text': ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml'],
            'document': ['.pdf', '.doc', '.docx'],
            'image': ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'],
            'spreadsheet': ['.xls', '.xlsx', '.csv']
        }
    
    async def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        處理文檔並提取內容
        
        Args:
            file_path: 文件路徑
            filename: 文件名
            
        Returns:
            處理結果字典
        """
        try:
            file_ext = Path(filename).suffix.lower()
            
            # 根據文件類型選擇處理方法
            if file_ext in self.supported_formats['text']:
                return await self._process_text_file(file_path, filename)
            elif file_ext in self.supported_formats['document']:
                return await self._process_document_file(file_path, filename)
            elif file_ext in self.supported_formats['image']:
                return await self._process_image_file(file_path, filename)
            elif file_ext in self.supported_formats['spreadsheet']:
                return await self._process_spreadsheet_file(file_path, filename)
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件類型: {file_ext}",
                    "supported_formats": self.supported_formats
                }
                
        except Exception as e:
            self.logger.error(f"文檔處理失敗: {e}")
            return {
                "success": False,
                "error": f"文檔處理失敗: {str(e)}"
            }
    
    async def _process_text_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """處理純文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 分析文本內容
            analysis = self._analyze_text_content(content, filename)
            
            return {
                "success": True,
                "type": "text",
                "filename": filename,
                "content": content,
                "analysis": analysis,
                "metadata": {
                    "file_size": os.path.getsize(file_path),
                    "line_count": len(content.splitlines()),
                    "char_count": len(content),
                    "word_count": len(content.split())
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"文本文件處理失敗: {str(e)}"
            }
    
    async def _process_document_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """處理文檔文件（PDF, DOC, DOCX）"""
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                return await self._process_pdf(file_path, filename)
            elif file_ext in ['.doc', '.docx']:
                return await self._process_word_document(file_path, filename)
            else:
                return {
                    "success": False,
                    "error": f"不支持的文檔類型: {file_ext}"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"文檔處理失敗: {str(e)}"
            }
    
    async def _process_pdf(self, file_path: str, filename: str) -> Dict[str, Any]:
        """處理PDF文件"""
        try:
            doc = fitz.open(file_path)
            
            # 提取文本內容
            text_content = ""
            images = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 提取文本
                page_text = page.get_text()
                text_content += f"\\n--- 第 {page_num + 1} 頁 ---\\n{page_text}\\n"
                
                # 提取圖像
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n - pix.alpha < 4:  # 確保是RGB或灰度圖像
                            img_data = pix.tobytes("png")
                            img_base64 = base64.b64encode(img_data).decode()
                            images.append({
                                "page": page_num + 1,
                                "index": img_index,
                                "data": img_base64,
                                "format": "png"
                            })
                        pix = None
                    except Exception as e:
                        self.logger.warning(f"提取PDF圖像失敗: {e}")
            
            doc.close()
            
            # 分析內容
            analysis = self._analyze_text_content(text_content, filename)
            
            return {
                "success": True,
                "type": "pdf",
                "filename": filename,
                "content": text_content,
                "images": images,
                "analysis": analysis,
                "metadata": {
                    "page_count": len(doc),
                    "image_count": len(images),
                    "file_size": os.path.getsize(file_path)
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF處理失敗: {str(e)}"
            }
    
    async def _process_word_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """處理Word文檔"""
        try:
            doc = Document(file_path)
            
            # 提取文本內容
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\\n"
            
            # 提取表格內容
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            # 分析內容
            analysis = self._analyze_text_content(text_content, filename)
            
            return {
                "success": True,
                "type": "word",
                "filename": filename,
                "content": text_content,
                "tables": tables_content,
                "analysis": analysis,
                "metadata": {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "file_size": os.path.getsize(file_path)
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Word文檔處理失敗: {str(e)}"
            }
    
    async def _process_image_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """處理圖像文件"""
        try:
            # 打開圖像
            with Image.open(file_path) as img:
                # 轉換為RGB（如果需要）
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # 獲取圖像信息
                width, height = img.size
                
                # 轉換為base64
                import io
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode()
            
            # 使用Cloud Search MCP進行OCR（如果可用）
            ocr_result = await self._perform_ocr(file_path, filename)
            
            return {
                "success": True,
                "type": "image",
                "filename": filename,
                "image_data": img_base64,
                "ocr_result": ocr_result,
                "metadata": {
                    "width": width,
                    "height": height,
                    "format": img.format,
                    "file_size": os.path.getsize(file_path)
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"圖像處理失敗: {str(e)}"
            }
    
    async def _process_spreadsheet_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """處理電子表格文件"""
        try:
            import pandas as pd
            
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.csv':
                df = pd.read_csv(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                df = pd.read_excel(file_path)
            else:
                return {
                    "success": False,
                    "error": f"不支持的電子表格類型: {file_ext}"
                }
            
            # 轉換為字典格式
            data = df.to_dict('records')
            
            # 生成摘要
            summary = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": list(df.columns),
                "data_types": df.dtypes.to_dict()
            }
            
            return {
                "success": True,
                "type": "spreadsheet",
                "filename": filename,
                "data": data,
                "summary": summary,
                "metadata": {
                    "file_size": os.path.getsize(file_path)
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"電子表格處理失敗: {str(e)}"
            }
    
    async def _perform_ocr(self, file_path: str, filename: str) -> Dict[str, Any]:
        """執行OCR識別"""
        try:
            # 嘗試使用Cloud Search MCP
            try:
                from aicore0619.mcp.adapter.cloud_search_mcp.cloud_search_mcp import CloudSearchMCP
                
                cloud_search = CloudSearchMCP()
                
                # 讀取圖像並轉換為base64
                with open(file_path, 'rb') as f:
                    img_data = f.read()
                    img_base64 = base64.b64encode(img_data).decode()
                
                # 執行OCR
                ocr_request = {
                    "image_data": img_base64,
                    "task_type": "ocr",
                    "language": "auto"
                }
                
                result = await cloud_search.process_vision_request(ocr_request)
                
                return {
                    "success": True,
                    "method": "cloud_search_mcp",
                    "content": result.get("content", ""),
                    "confidence": result.get("confidence", 0.0)
                }
                
            except Exception as e:
                self.logger.warning(f"Cloud Search MCP OCR失敗: {e}")
                
                # 使用本地OCR作為備用
                return await self._local_ocr_fallback(file_path)
                
        except Exception as e:
            return {
                "success": False,
                "error": f"OCR處理失敗: {str(e)}"
            }
    
    async def _local_ocr_fallback(self, file_path: str) -> Dict[str, Any]:
        """本地OCR備用方案"""
        try:
            # 嘗試使用pytesseract
            try:
                import pytesseract
                from PIL import Image
                
                img = Image.open(file_path)
                text = pytesseract.image_to_string(img, lang='chi_tra+chi_sim+eng')
                
                return {
                    "success": True,
                    "method": "pytesseract",
                    "content": text,
                    "confidence": 0.7  # 估計置信度
                }
                
            except ImportError:
                self.logger.warning("pytesseract未安裝，跳過本地OCR")
                
            # 簡單的圖像描述
            return {
                "success": True,
                "method": "basic_analysis",
                "content": "圖像文件已上傳，但無法進行文字識別。請確保安裝了OCR工具。",
                "confidence": 0.0
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"本地OCR失敗: {str(e)}"
            }
    
    def _analyze_text_content(self, content: str, filename: str) -> Dict[str, Any]:
        """分析文本內容"""
        try:
            # 基本統計
            lines = content.splitlines()
            words = content.split()
            
            # 檢測可能的需求關鍵詞
            requirement_keywords = [
                '需求', '要求', '功能', '特性', '規格', '設計', '開發', '實現',
                'requirement', 'feature', 'function', 'specification', 'design',
                'develop', 'implement', '用戶故事', 'user story', '驗收標準'
            ]
            
            found_keywords = []
            for keyword in requirement_keywords:
                if keyword.lower() in content.lower():
                    found_keywords.append(keyword)
            
            # 檢測結構化內容
            has_lists = any(line.strip().startswith(('-', '*', '•', '1.', '2.')) for line in lines)
            has_headers = any(line.strip().startswith('#') for line in lines)
            
            return {
                "word_count": len(words),
                "line_count": len(lines),
                "char_count": len(content),
                "requirement_keywords": found_keywords,
                "has_structured_content": has_lists or has_headers,
                "has_lists": has_lists,
                "has_headers": has_headers,
                "estimated_reading_time": max(1, len(words) // 200),  # 分鐘
                "content_type": self._detect_content_type(content, filename)
            }
            
        except Exception as e:
            self.logger.error(f"內容分析失敗: {e}")
            return {
                "error": f"內容分析失敗: {str(e)}"
            }
    
    def _detect_content_type(self, content: str, filename: str) -> str:
        """檢測內容類型"""
        content_lower = content.lower()
        
        # 檢測代碼
        code_indicators = ['def ', 'function ', 'class ', 'import ', 'from ', '<?php', '<html', 'SELECT ', 'CREATE TABLE']
        if any(indicator in content for indicator in code_indicators):
            return "code"
        
        # 檢測需求文檔
        requirement_indicators = ['需求分析', '功能需求', '非功能需求', '用戶故事', 'user story', 'requirement']
        if any(indicator in content_lower for indicator in requirement_indicators):
            return "requirement_document"
        
        # 檢測技術文檔
        tech_indicators = ['api', '接口', '架構', 'architecture', '設計', 'design', '技術方案']
        if any(indicator in content_lower for indicator in tech_indicators):
            return "technical_document"
        
        # 檢測配置文件
        if filename.endswith(('.json', '.yaml', '.yml', '.toml', '.ini', '.conf')):
            return "configuration"
        
        return "general_text"

# 全局實例
document_processor = MultimodalDocumentProcessor()

